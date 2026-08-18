import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def set_font(run, size, bold=False):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold

def add_paragraph(doc, text, size=12, bold=False, align=WD_PARAGRAPH_ALIGNMENT.LEFT, space_after=Pt(12)):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = space_after
    p.alignment = align
    run = p.add_run(text)
    set_font(run, size, bold)
    return p

def main():
    doc = Document()
    
    # Set default style to Times New Roman 12
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # ==========================================
    # Title Page
    # ==========================================
    add_paragraph(doc, "2026", 16, True, WD_PARAGRAPH_ALIGNMENT.CENTER, Pt(0))
    add_paragraph(doc, "International Internship Pilot Program", 16, True, WD_PARAGRAPH_ALIGNMENT.CENTER, Pt(0))
    add_paragraph(doc, "in Taiwan", 16, True, WD_PARAGRAPH_ALIGNMENT.CENTER, Pt(24))
    
    add_paragraph(doc, "Final Report", 18, True, WD_PARAGRAPH_ALIGNMENT.CENTER, Pt(48))
    
    add_paragraph(doc, "Name : Yadav/Rikesh", 12, False)
    add_paragraph(doc, "Internship Inst. : Yuan Ze University, Taoyuan, Taiwan", 12, False)
    add_paragraph(doc, "Name of PI/Department/School : Asst. Prof. Dr. Ihsan Ullah / INEC Laboratory / Computer Science Engineering", 12, False)
    add_paragraph(doc, "Submission Date : 2026/08/29", 12, False)
    
    doc.add_page_break()

    # ==========================================
    # Part I. Overview
    # ==========================================
    add_paragraph(doc, "Part I. Overview", 16, True, WD_PARAGRAPH_ALIGNMENT.CENTER, Pt(24))

    q1 = doc.add_paragraph()
    q1.paragraph_format.line_spacing = 1.5
    set_font(q1.add_run("◼ What is your intention to come to Taiwan for this internship?"), 12, True)
    add_paragraph(doc, "My primary intention to come to Taiwan was to gain exposure to a leading international research environment and to immerse myself in a new academic culture. Taiwan is globally recognized for its advancements in technology and edge computing, making the Intelligent Networks and Edge-Cloud Computing (INEC) Laboratory at Yuan Ze University the perfect place to advance my understanding of Deep Reinforcement Learning (DRL) and Graph Neural Networks (GNNs). Furthermore, I sought an independent research environment that would challenge me to develop functional prototypes geared toward prestigious IEEE conference publications.")

    q2 = doc.add_paragraph()
    q2.paragraph_format.line_spacing = 1.5
    set_font(q2.add_run("◼ Please describe your gains and experience acquired from this internship."), 12, True)
    add_paragraph(doc, "During this 12-week structured internship, I successfully acquired profound theoretical and practical expertise. Starting from fundamental Q-Learning in custom Python GridWorlds, I progressed through Deep Q-Networks (DQN) applied to CartPole and LunarLander environments, eventually mastering PyTorch Geometric to build Graph Convolutional Networks (GCNs). My most significant gain was bridging these two disparate fields to engineer a novel, unified GCN-DQN architecture capable of real-time, adaptive decision-making for IoT-enabled smart environment evacuation.")

    q3 = doc.add_paragraph()
    q3.paragraph_format.line_spacing = 1.5
    set_font(q3.add_run("◼ Please describe your life in Taiwan, how to deal with difficulties?"), 12, True)
    add_paragraph(doc, "Living and studying in Taiwan has been an incredibly meaningful personal experience. Initially, adapting to the independent research environment and the fast-paced academic culture presented a learning curve. To overcome these difficulties, I maintained regular communication with my Principal Investigator through weekly progress meetings. I also actively interacted with fellow IIPP interns, creating a genuinely collaborative support system that helped me navigate both academic hurdles and daily life in a new country.")

    q4 = doc.add_paragraph()
    q4.paragraph_format.line_spacing = 1.5
    set_font(q4.add_run("◼ During the internship, what impressed you most and made you recommend the internship to your friends?"), 12, True)
    add_paragraph(doc, "What impressed me the most was the perfect balance between guided supervision and research autonomy provided by the IIPP and the INEC Lab. Rather than just executing predefined tasks, I was encouraged to conduct focused literature reviews on 2024 and 2025 IEEE publications, identify research gaps, and formulate my own problem statement. The access to state-of-the-art computational resources and the warm, collaborative international community at Yuan Ze University makes this an experience I will highly recommend to my peers.")

    q5 = doc.add_paragraph()
    q5.paragraph_format.line_spacing = 1.5
    set_font(q5.add_run("◼ If there is any other similar chance for you to choose, would you like to participate in this program again? And why?"), 12, True)
    add_paragraph(doc, "Yes, absolutely. Participating in this program again would allow me to further extend my current research—such as expanding the single-agent evacuation framework into a Multi-Agent Reinforcement Learning (MARL) paradigm. The exposure to Taiwan’s rich academic ecosystem is invaluable, and I look forward to carrying this international collaboration into my future graduate studies and beyond.")

    q6 = doc.add_paragraph()
    q6.paragraph_format.line_spacing = 1.5
    set_font(q6.add_run("◼ Please describe during your internship, which PI and institute have you communicated or interacted with?"), 12, True)
    add_paragraph(doc, "I have maintained consistent, direct communication with Assistant Professor Dr. Ihsan Ullah at the Intelligent Networks and Edge-Cloud Computing (INEC) Laboratory, Yuan Ze University. Through weekly meetings, we discussed study outcomes, implementation results, and strategic research planning that guided my project to its successful conclusion.")

    doc.add_page_break()

    # ==========================================
    # Part II. Research Finding
    # ==========================================
    add_paragraph(doc, "Part II. Research Finding", 16, True, WD_PARAGRAPH_ALIGNMENT.CENTER, Pt(24))
    
    add_paragraph(doc, "Intelligent IoT Sensor-Based Dynamic Evacuation Framework using Hybrid GCN-DQN", 18, True, WD_PARAGRAPH_ALIGNMENT.CENTER, Pt(24))

    add_paragraph(doc, "I. Introduction", 16, True)
    add_paragraph(doc, "Existing Graph Neural Network (GNN) and Deep Reinforcement Learning (DRL) approaches either treat the two as separate modules or focus exclusively on vehicular environments. No prior work has proposed a unified GCN-DQN architecture for intelligent decision-making in stationary IoT environments. After in-depth discussions with Dr. Ihsan Ullah and evaluating existing literature, the research focus was finalized on developing a dynamic evacuation framework for smart environments. This system combines graph-based spatial representations with reinforcement learning to make adaptive evacuation decisions using real-time sensor data.")

    add_paragraph(doc, "II. Deep Reinforcement Learning Foundations (Weeks 1-4)", 16, True)
    add_paragraph(doc, "A. Core RL Implementations", 14, True)
    add_paragraph(doc, "The first phase of research was dedicated to building a strong DRL foundation. I implemented the core Markov Decision Process (MDP) framework, constructing a GridWorld environment entirely from scratch without external libraries. The agent successfully learned an optimal navigation policy through iterative Q-table updates governed by the Bellman equation. Subsequently, I transitioned to PyTorch to develop Deep Q-Networks (DQN), successfully solving both the CartPole-v1 and LunarLander-v2 continuous-state environments.")
    add_paragraph(doc, "B. Algorithmic Enhancements", 14, True)
    add_paragraph(doc, "A key innovation that stabilized the early DRL implementation was the introduction of the Experience Replay Buffer, which stores past transitions and samples random mini-batches to break temporal correlations. Further literature review during Week 3 and 4 covered Double DQN (reducing overestimation bias), Dueling DQN, and Prioritized Experience Replay.")

    add_paragraph(doc, "III. Graph Neural Networks & Spatial Modeling (Weeks 5-8)", 16, True)
    add_paragraph(doc, "A. Transitioning from Grids to Graphs", 14, True)
    add_paragraph(doc, "Week 5 marked a meaningful shift toward Graph Convolutional Networks (GCNs). Traditional Convolutional Neural Networks struggle to generalize across varying building sizes. By representing buildings as graphs—where nodes are physical locations and edges are valid walking paths—the environment becomes scale-invariant. Using PyTorch Geometric (PyG), I developed automated algorithms to map dynamic 2D spatial layouts into PyG Data objects.")
    add_paragraph(doc, "B. Node Classification on Custom Graphs", 14, True)
    add_paragraph(doc, "Initial GCN experiments involved applying the Kipf & Welling GCN layer formula to classify nodes in the standard Cora citation dataset, followed by a custom-built IoT edge-network topology to predict node congestion states (Idle vs. Busy) using masked transductive learning.")

    add_paragraph(doc, "IV. The Hybrid GCN-DQN Architecture (Weeks 9-11)", 16, True)
    add_paragraph(doc, "A. Unified Model Design", 14, True)
    add_paragraph(doc, "The culmination of the theoretical research was the design of a hybrid GNNDQNetwork class. The GCN acts as the spatial feature extractor, processing dynamic hazards like fire and smoke across the building graph. To interface this with the DQN action head, I implemented a novel 'Agent-Centric Pooling' mechanism. Instead of aggregating the entire graph via global mean pooling, this mechanism isolates the specific 64-dimensional hidden vector corresponding exclusively to the node where the agent is currently located.")
    add_paragraph(doc, "B. Hazard-Aware Reward Shaping", 14, True)
    add_paragraph(doc, "To force the agent away from danger, I integrated Hazard-Aware Dense Reward shaping. Active fire nodes are dynamically masked as impassable entities in the PyG graph, causing the A* potential heuristic to spike, instantly repelling the DQN from dangerous corridors.")

    add_paragraph(doc, "V. Experimental Setup and Results", 16, True)
    add_paragraph(doc, "A. Zero-Shot Transferability", 14, True)
    add_paragraph(doc, "The GCN-DQN agent was trained strictly on an Office (10x10) environment for 1500 episodes with randomized fire spreading. Following training, the model's weights were frozen, and it was deployed sequentially into massive, unseen layouts. Because the model learned the topological relationship of hazards rather than absolute coordinates, it achieved a near 100% successful evacuation rate zero-shot in the sprawling 30x30 Mall environment.")
    add_paragraph(doc, "B. IoT Telemetry Dashboard", 14, True)
    add_paragraph(doc, "To simulate real-world IoT deployment, the GNN-DQN inference engine was decoupled via a FastAPI backend, communicating with a premium React-based dashboard via WebSockets. This successfully demonstrated real-time, edge-processed dynamic routing for smart building occupants.")

    add_paragraph(doc, "VI. Conclusion and Future Work", 16, True)
    add_paragraph(doc, "The developed GCN-DQN framework successfully overcomes the generalization failures of standard grid-based DRL, offering a highly robust prototype for intelligent environments. Future research will expand this framework into the Multi-Agent Reinforcement Learning (MARL) domain, targeting cooperative bottleneck resolution and crowd density management during mass evacuations.")

    # Save the document
    output_path = os.path.abspath('Final_Report_Rikesh_Yadav.docx')
    doc.save(output_path)
    print(f"Report successfully generated at: {output_path}")

if __name__ == '__main__':
    main()
